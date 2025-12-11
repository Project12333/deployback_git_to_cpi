#!/usr/bin/env python3
"""
Generate CPI iFlow Technical Specification DOCX:
- Uses DeepSeek (Ollama) to produce structured Markdown
- Converts Markdown -> DOCX using pandoc (with TOC)
- Injects SAP + Motiveminds logos into header (every page)
- Adds right-aligned page numbers in footer
- No reference.docx required
"""

import sys
import json
import requests
import xml.etree.ElementTree as ET
from pathlib import Path
import subprocess
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# -------------------------
# Paths & config
# -------------------------
BASE_DIR = Path(__file__).resolve().parent            # tools/
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434/api/generate")
MODEL_NAME = os.getenv("OLLAMA_MODEL", "deepseek-r1:1.5b")

SAP_LOGO = BASE_DIR / "logos" / "sap.png"
MM_LOGO = BASE_DIR / "logos" / "motiveminds.png"

# -------------------------
# Helpers: parse .iflw
# -------------------------
def parse_iflw(path):
    meta = {
        "flowname": Path(path).stem,
        "senders": [],
        "receivers": [],
        "adapters": [],
        "scripts": [],
        "mappings": [],
        "steps": []
    }
    try:
        tree = ET.parse(path)
        root = tree.getroot()
        for elem in root.iter():
            tag = elem.tag.split("}")[-1].lower()
            name = elem.attrib.get("name") or elem.attrib.get("id") or elem.attrib.get("idref") or tag

            if "sender" in tag or "start" in tag:
                if name not in meta["senders"]:
                    meta["senders"].append(name)
            if "receiver" in tag or "end" in tag:
                if name not in meta["receivers"]:
                    meta["receivers"].append(name)
            if "adapter" in tag:
                if name not in meta["adapters"]:
                    meta["adapters"].append(name)
            if "script" in tag or "groovy" in tag or "scripttask" in tag:
                if name not in meta["scripts"]:
                    meta["scripts"].append(name)
            if "mapping" in tag:
                if name not in meta["mappings"]:
                    meta["mappings"].append(name)
            if name not in meta["steps"]:
                meta["steps"].append(name)
    except Exception as e:
        meta["error"] = f"XML parse error: {e}"
    return meta

# -------------------------
# Mermaid diagram
# -------------------------
def sanitize_id(s):
    return "".join(c if c.isalnum() else "_" for c in str(s))

def high_level_diagram(meta):
    sender = meta["senders"][0] if meta["senders"] else "Sender"
    receiver = meta["receivers"][0] if meta["receivers"] else "Receiver"
    return (
        "```mermaid\n"
        "graph TD\n"
        f"  {sanitize_id(sender)}([\"{sender}\"]) -->|Request| CPI\n"
        f"  CPI -->|Response| {sanitize_id(receiver)}([\"{receiver}\"])\n"
        "```\n"
    )

# -------------------------
# Prompt builder (strict)
# -------------------------
def build_prompt(meta, diagram):
    # Strict prompt: produces ONLY markdown document with TOC marker
    return (
        "You are a senior SAP CPI integration architect and technical writer. "
        "Produce ONE complete Technical Specification document in MARKDOWN only. "
        "Do NOT output any explanations, analysis, or extra text — only the Markdown document.\n\n"

        "<h1 style=\"color: #1f4e79; font-size: 2.2em;\">Table of Contents</h1>\n\n"
        "1. Introduction\n"
        "  1.1 Purpose\n"
        "  1.2 Scope\n"
        "2. Integration Overview\n"
        "  2.1 Integration Architecture\n"
        "  2.2 Integration Components\n"
        "3. Integration Scenarios\n"
        "  3.1 Scenario Description\n"
        "  3.2 Data Flows\n"
        "  3.3 Security Requirements\n"
        "4. Error Handling and Logging\n"
        "5. Testing Validation\n"
        "6. Reference Documents\n\n"

        "---TOC-END-PAGE-BREAK---\n\n"

        "# 1. Introduction\n\n"
        "## 1.1 Purpose\n"
        "Write a concise purpose for this iFlow using the metadata below. If metadata is missing, state a single-line suggestion of what to include.\n\n"
        "## 1.2 Scope\n"
        "Describe the scope and boundaries of the iFlow (systems affected, data types, what is out of scope).\n\n"

        "# 2. Integration Overview\n\n"
        "## 2.1 Integration Architecture\n"
        "Provide a concise architecture description referencing the sender(s), CPI, and receiver(s). Then insert the High-Level Process Flow Diagram using Mermaid exactly as shown below.\n\n"
        + diagram + "\n\n"  # ensure one blank line after mermaid fence

        "## 2.2 Integration Components\n"
        "Provide a short table/list of sender/receiver systems and adapter types. Use the metadata below to fill the values.\n\n"

        "# 3. Integration Scenarios\n\n"
        "## 3.1 Scenario Description\n"
        "Give a step-by-step (numbered) explanation of the iFlow processing path (message entry to exit).\n\n"
        "## 3.2 Data Flows\n"
        "Summarize mapping logic (XSLT / message mapping). If Groovy or other scripts exist, include small fenced code blocks with the script name and explain its purpose.\n\n"
        "## 3.3 Security Requirements\n"
        "List authentication methods, credentials usage, encryption, and any header/certificate requirements.\n\n"

        "# 4. Error Handling and Logging\n"
        "Describe exception subprocesses, retry logic, fallback actions, and what is logged (include sample log lines if applicable).\n\n"

        "# 5. Testing Validation\n"
        "List the key test cases to validate the iFlow (unit tests, E2E, negative tests, data transformation checks).\n\n"

        "# 6. Reference Documents\n"
        "List the input artifacts that were analyzed (iFlow file names, groovy scripts, XSLT, mapping sheets). Provide relative file paths when possible.\n\n"

        "### METADATA (use to fill sections, do NOT include this raw block in final doc):\n"
        + json.dumps(meta, indent=2)
        + "\n\n"

        "IMPORTANT INSTRUCTIONS:\n"
        "1) Output ONLY the Markdown document following the headings exactly as above.\n"
        "2) Use the flow name from metadata as the main visible title where appropriate.\n"
        "3) Keep each subsection concise but complete (prefer bullet lists and numbered steps where appropriate).\n"
        "4) When including code/script excerpts, wrap them in triple-backtick fenced blocks and include the file name as a small caption.\n"
        "5) If any metadata field is empty, place a single-line placeholder like '<MISSING: description of what to provide>'.\n"
        "6) Ensure the mermaid block has exactly one blank line after the closing ``` fence.\n"
        "7) Do NOT include any extraneous commentary, headers, or footers — pure Markdown only.\n"
    )

# -------------------------
# Call LLM
# -------------------------
def call_llm(prompt):
    payload = {"model": MODEL_NAME, "prompt": prompt, "stream": False}
    res = requests.post(OLLAMA_URL, json=payload, timeout=180)
    res.raise_for_status()
    j = res.json()
    # Ollama responses vary; try common keys
    for k in ("response", "text", "output"):
        if k in j and j[k]:
            return j[k]
    return json.dumps(j)

# -------------------------
# Markdown helpers
# -------------------------
def replace_toc_marker(md):
    # Pandoc will create a TOC with --toc; we still want a physical page break after TOC.
    # Replace our marker with an HTML page-break that pandoc will honor.
    return md.replace("---TOC-END-PAGE-BREAK---", "<div style=\"page-break-after: always;\"></div>")

# -------------------------
# Pandoc convert md -> docx with TOC
# -------------------------
def pandoc_md_to_docx(md_path: Path, docx_path: Path):
    cmd = [
        "pandoc",
        str(md_path),
        "--toc",
        "--toc-depth=2",
        "-o",
        str(docx_path)
    ]
    subprocess.run(cmd, check=True)

# -------------------------
# Insert header logos (every page) and right-aligned page numbers in footer
# -------------------------
def insert_header_logos_and_footer_pagenum(docx_path: Path, sap_logo: Path, mm_logo: Path):
    doc = Document(str(docx_path))

    # HEADER: clear and add 2-col table with logos
    header = doc.sections[0].header
    # Clear header paragraphs if any
    for p in header.paragraphs:
        p.clear()

    # add_table requires width parameter in some python-docx versions
    table = header.add_table(rows=1, cols=2, width=Inches(6))
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3)

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    # Left logo
    if sap_logo.exists():
        p_left = left_cell.paragraphs[0]
        rleft = p_left.add_run()
        rleft.add_picture(str(sap_logo), width=Inches(1.2))
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Right logo
    if mm_logo.exists():
        p_right = right_cell.paragraphs[0]
        rright = p_right.add_run()
        rright.add_picture(str(mm_logo), width=Inches(1.2))
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # FOOTER: add right-aligned page number field
    for section in doc.sections:
        footer = section.footer
        # Clear existing footer paragraphs
        for p in footer.paragraphs:
            p.clear()
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Insert page number field
        run = p.add_run()
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'PAGE')
        # Append field element into run
        run._r.append(fld)

    # Save back
    doc.save(str(docx_path))

# -------------------------
# Utilities
# -------------------------
def determine_version(flow_name):
    mode = os.getenv("VERSION_MODE", "none")
    manual = os.getenv("VERSION_VALUE", "")
    if mode == "manual" and manual:
        return manual
    if mode == "date":
        return datetime.now().strftime("%Y-%m-%d")
    if mode == "git-sha":
        try:
            sha = subprocess.check_output(["git", "rev-parse", "--short", "HEAD"]).decode().strip()
            return sha
        except:
            return "unknown"
    return "1.0"

# -------------------------
# Main writer
# -------------------------
def write_output(iflow_path: str, markdown: str, meta: dict):
    md_dir = Path(iflow_path).parent / "docs"
    md_dir.mkdir(parents=True, exist_ok=True)

    flow = meta.get("flowname", Path(iflow_path).stem)
    version = determine_version(flow)
    today = datetime.now().strftime("%Y-%m-%d")
    author = os.getenv("DOC_AUTHOR", "Sindhu K V")

    md_file = md_dir / f"{flow}.md"
    temp_docx = md_dir / f"{flow}_temp.docx"
    final_docx = md_dir / f"{flow}_Documentation_v{version}.docx"

    # Insert the flow title as top heading if not present
    if not markdown.strip().startswith("#"):
        markdown = f"# {flow}\n\n" + markdown

    # Ensure TOC page-break marker replaced
    markdown = replace_toc_marker(markdown)

    # Save md
    md_file.write_text(markdown, encoding="utf-8")

    # Convert to docx using pandoc with TOC
    pandoc_md_to_docx(md_file, temp_docx)

    # Now insert logos + page numbers into the pandoc-generated docx
    insert_header_logos_and_footer_pagenum(temp_docx, SAP_LOGO, MM_LOGO)

    # Move/rename to final file
    temp_docx.replace(final_docx)

    print(f"✔ Generated final DOCX: {final_docx}")
    return final_docx

# -------------------------
# MAIN
# -------------------------
def main(argv):
    if len(argv) < 2:
        print("Usage: ollama_generate_docs.py <iflow.iflw> [more.iflw ...]")
        sys.exit(1)

    for path in argv[1:]:
        print(f"Processing {path}")
        meta = parse_iflw(path)
        diagram = high_level_diagram(meta)
        prompt = build_prompt(meta, diagram)

        print("Calling DeepSeek (Ollama)...")
        try:
            markdown = call_llm(prompt)
        except Exception as e:
            print("LLM call failed:", e)
            markdown = "# " + meta.get("flowname", Path(path).stem) + "\n\n<Failed to generate text>"

        # Ensure the flow name is heading 1 at top if model didn't add it
        if not markdown.lstrip().startswith("#"):
            markdown = f"# {meta.get('flowname', Path(path).stem)}\n\n" + markdown

        # Write output docx with logos and page numbers
        try:
            write_output(path, markdown, meta)
        except Exception as e:
            print("Failed to write output:", e)
            raise

if __name__ == "__main__":
    main(sys.argv)
