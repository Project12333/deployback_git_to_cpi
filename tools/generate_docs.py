import os
import argparse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from tools.ollama_client import generate_section

SAP_LOGO = "tools/logos/sap.png"
MOTIVE_LOGO = "tools/logos/motiveminds.png"

# -------------------------------------------------
# Header with logos (ALL pages)
# -------------------------------------------------
def add_header_with_logos(section):
    header = section.header
    header.is_linked_to_previous = False

    table = header.add_table(rows=1, cols=2)

    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left.add_run().add_picture(SAP_LOGO, width=Inches(1.2))

    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run().add_picture(MOTIVE_LOGO, width=Inches(1.5))

# -------------------------------------------------
# Word Table of Contents
# -------------------------------------------------
def add_table_of_contents(doc):
    p = doc.add_paragraph()
    r = p.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(ns.qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    end = OxmlElement("w:fldChar")
    end.set(ns.qn("w:fldCharType"), "end")

    r._r.extend([begin, instr, end])

# -------------------------------------------------
# Main generator
# -------------------------------------------------
def generate_doc(package_name):
    output_dir = f"cpi-artifacts/{package_name}/docs"
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    add_header_with_logos(section)

    # ---------------- Cover Page ----------------
    title = doc.add_heading(
        "SAP CPI Integration Flow\nTechnical Specification",
        level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ---------------- TOC Page ----------------
    doc.add_heading("Table of Contents", level=1)
    add_table_of_contents(doc)
    doc.add_page_break()

    # ---------------- TOC-driven sections ----------------
    sections = [
        ("1. Introduction", "Overview of the SAP CPI integration flow"),
        ("1.1 Purpose", "Purpose of this integration"),
        ("1.2 Scope", "Scope and limitations"),
        ("2. Integration Overview", "High-level overview"),
        ("2.1 Integration Architecture", "Architecture of the CPI flow"),
        ("2.2 Integration Components", "Adapters, mappings, scripts"),
        ("3. Integration Scenarios", "Business scenarios supported"),
        ("3.1 Scenario Description", "Detailed scenario explanation"),
        ("3.2 Data Flows", "Inbound and outbound data flows"),
        ("3.3 Security Requirements", "Authentication and authorization"),
        ("4. Error Handling and Logging", "Exception handling and logging"),
        ("5. Testing Validation", "Testing and validation approach"),
        ("6. Reference Documents", "Related SAP documentation"),
    ]

    for title, context in sections:
        print(f"🧠 Generating section: {title}")  # 🔑 progress indicator
        level = 1 if title.count(".") == 1 else 2
        doc.add_heading(title, level=level)
        doc.add_paragraph(generate_section(title, context))

    output_file = f"{output_dir}/{package_name}_Technical_Spec.docx"
    doc.save(output_file)
    print(f"✅ Document generated: {output_file}")
    print("ℹ Open Word → Right-click TOC → Update Field")

# -------------------------------------------------
# Entry point
# -------------------------------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate SAP CPI documentation")
    parser.add_argument("--package", required=True, help="CPI package name")
    args = parser.parse_args()

    generate_doc(args.package)

