from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import sys
import os

# ---------------- ARGUMENTS ----------------
# 1 = input text (AI-generated content)
# 2 = output docx path
# 3 = iFlow name
# 4 = author
# 5 = version
input_text = sys.argv[1]
output_docx = sys.argv[2]
iflow_name = sys.argv[3]
author = sys.argv[4]
version = sys.argv[5]

# ---------------- DOCUMENT ----------------
doc = Document()

# ---------------- HEADER (LOGOS ON ALL PAGES) ----------------
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

# SAP Logo (Left)
run_left = header_para.add_run()
run_left.add_picture("tools/logos/sap.png", width=Inches(1.6))

# Spacer
header_para.add_run("\t" * 6)

# MotiveMinds Logo (Right)
run_right = header_para.add_run()
run_right.add_picture("tools/logos/motiveminds.png", width=Inches(1.5))

# ---------------- COVER PAGE ----------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run(iflow_name)
title_run.bold = True
title_run.font.size = Inches(0.35)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Technical Specification Document")

doc.add_paragraph("\n")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Author: {author}\n")
meta.add_run(f"Date: {datetime.today().strftime('%Y-%m-%d')}\n")
meta.add_run(f"Version: {version}")

doc.add_page_break()

# ---------------- MAIN CONTENT ----------------
with open(input_text, "r", encoding="utf-8") as f:
    for line in f:
        doc.add_paragraph(line.rstrip())

# ---------------- SAVE ----------------
os.makedirs(os.path.dirname(output_docx), exist_ok=True)
doc.save(output_docx)

print(f"Document generated successfully: {output_docx}")
