import argparse
import os
from pathlib import Path
import xml.etree.ElementTree as ET
from datetime import date

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ollama_client import call_ollama

# -------------------------------------------------
# Arguments
# -------------------------------------------------
parser = argparse.ArgumentParser(description="Generate SAP CPI FS document")
parser.add_argument("--package", required=True)
parser.add_argument("--author", default="Sindhu")
args = parser.parse_args()

PACKAGE = args.package
AUTHOR = args.author

print("🚀 Script started")

# -------------------------------------------------
# Paths
# -------------------------------------------------
BASE_DIR = Path.cwd()
PACKAGE_DIR = BASE_DIR / "cpi-artifacts" / PACKAGE

print("📦 Package path:", PACKAGE_DIR.resolve())

if not PACKAGE_DIR.exists():
    raise SystemExit("❌ Package not found")

OUTPUT_DIR = BASE_DIR / "generated-docs" / PACKAGE
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

LOGOS_DIR = BASE_DIR / "tools" / "logos"
SAP_LOGO = LOGOS_DIR / "sap.png"
MM_LOGO = LOGOS_DIR / "motiveminds.png"

# -------------------------------------------------
# Header (logos)
# -------------------------------------------------
def add_header(doc):
    header = doc.sections[0].header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False

    left = table.cell(0, 0).paragraphs[0]
    left.add_run().add_picture(str(SAP_LOGO), width=Inches(1.2))

    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run().add_picture(str(MM_LOGO), width=Inches(1.5))

# -------------------------------------------------
# Cover Page
# -------------------------------------------------
def add_cover_page(doc, flow_name):
    doc.add_heading(flow_name, level=0)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"

    table.cell(0, 0).text = "Author"
    table.cell(0, 1).text = AUTHOR

    table.cell(1, 0).text = "Date"
    table.cell(1, 1).text = date.today().isoformat()

    table.cell(2, 0).text = "Version"
    table.cell(2, 1).text = "Draft"

    doc.add_page_break()

# -------------------------------------------------
# Table of Contents (static FS)
# -------------------------------------------------
def add_toc(doc):
    doc.add_heading("Table of Contents", level=1)

    toc = [
        "1. Introduction",
        "   1.1 Purpose",
        "   1.2 Scope",
        "",
        "2. Integration Overview",
        "   2.1 Integration Architecture",
        "   2.2 Integration Components",
        "",
        "3. Integration Scenarios",
        "   3.1 Scenario Description",
        "   3.2 Data Flows",
        "   3.3 Security Requirements",
        "",
        "4. Error Handling and Logging",
        "5. Testing Validation",
        "6. Reference Documents",
    ]

    for line in toc:
        doc.add_paragraph(line)

    doc.add_page_break()

# -------------------------------------------------
# Parse iFlow XML
# -------------------------------------------------
def parse_iflow(iflw_path):
    components = set()
    try:
        root = ET.parse(iflw_path).getroot()
        for elem in root.iter():
            if "name" in elem.attrib:
                components.add(elem.attrib["name"])
    except Exception:
        pass

    return {
        "flow_name": iflw_path.stem,
        "components": ", ".join(sorted(components)) if components else "Not available"
    }

# -------------------------------------------------
# Ollama Prompt (STRICT FS MODE)
# -------------------------------------------------
def build_prompt(meta):
    return f"""
You are writing content for a SAP CPI Functional Specification.

STRICT RULES:
- Do NOT use markdown
- Do NOT use bold text
- Do NOT add headings or numbering
- Do NOT summarize sections
- Do NOT write introductions like "this document outlines"

Write ONLY implementation-level content under each label.

Purpose:
Scope:
Integration Architecture:
Integration Components:
Scenario Description:
Data Flows:
Security Requirements:
Error Handling and Logging:
Testing Validation:
Reference Documents:

Context:
Package: {PACKAGE}
Integration Flow: {meta['flow_name']}
Components: {meta['components']}
"""

# -------------------------------------------------
# Extract section text
# -------------------------------------------------
def extract_section(text, label):
    if label not in text:
        return ""

    start = text.index(label) + len(label)
    end = len(text)

    for other in [
        "Purpose:", "Scope:", "Integration Architecture:",
        "Integration Components:", "Scenario Description:",
        "Data Flows:", "Security Requirements:",
        "Error Handling and Logging:", "Testing Validation:",
        "Reference Documents:"
    ]:
        if other != label and other in text[start:]:
            end = min(end, text.find(other, start))

    return text[start:end].strip()

# -------------------------------------------------
# Discover iFlows (Windows safe)
# -------------------------------------------------
iflows = []
for root, dirs, files in os.walk(PACKAGE_DIR):
    for f in files:
        if f.lower().endswith(".iflw"):
            iflows.append(Path(root) / f)

print(f"🔍 iFlows found: {len(iflows)}")

if not iflows:
    raise SystemExit("⚠ No .iflw files found")

# -------------------------------------------------
# Generate document
# -------------------------------------------------
for iflw in iflows:
    print(f"\n➡ Processing iFlow: {iflw.stem}")

    meta = parse_iflow(iflw)

    print("🤖 Calling Ollama...")
    ai_text = call_ollama(build_prompt(meta))

    doc = Document()
    add_header(doc)
    add_cover_page(doc, meta["flow_name"])
    add_toc(doc)

    sections = [
        ("1. Introduction", None),
        ("1.1 Purpose", "Purpose:"),
        ("1.2 Scope", "Scope:"),
        ("2. Integration Overview", None),
        ("2.1 Integration Architecture", "Integration Architecture:"),
        ("2.2 Integration Components", "Integration Components:"),
        ("3. Integration Scenarios", None),
        ("3.1 Scenario Description", "Scenario Description:"),
        ("3.2 Data Flows", "Data Flows:"),
        ("3.3 Security Requirements", "Security Requirements:"),
        ("4. Error Handling and Logging", "Error Handling and Logging:"),
        ("5. Testing Validation", "Testing Validation:"),
        ("6. Reference Documents", "Reference Documents:")
    ]

    for title, label in sections:
        doc.add_heading(title, level=1 if title.count(".") == 1 else 2)

        if label:
            content = extract_section(ai_text, label)
            for line in content.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())

    output = OUTPUT_DIR / f"{meta['flow_name']}.docx"
    doc.save(output)

    print(f"✅ Generated: {output}")

print("\n🎉 Documentation generation completed successfully")
