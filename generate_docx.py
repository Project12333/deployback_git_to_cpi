import sys, json, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_numbered_heading(doc, text, level, num_str):
    p = doc.add_heading(level=level)
    run = p.add_run(f"{num_str} {text}")
    run.font.color.rgb = RGBColor(31, 56, 100)
    run.font.size = Pt(16 if level == 1 else 13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def build_docx(iflow_name, author, date, json_raw, output_path):
    doc = Document()
    try:
        data = json.loads(json_raw)
    except:
        data = {}

    # --- HEADER (LOCAL LOGOS - INCREASED SIZE) ---
    section = doc.sections[0]
    header = section.header
    htable = header.add_table(1, 2, width=Inches(6.5))
    htable.allow_autofit = False
    
    # Using your local repository paths for high-quality "bright" logos
    sap_path = "tools/logos/sap.png"
    mm_path = "tools/logos/motiveminds.png"

    if os.path.exists(sap_path):
        p1 = htable.rows[0].cells[0].paragraphs[0]
        # Increased height to 0.6 for better visibility
        p1.add_run().add_picture(sap_path, height=Inches(0.6))
    
    if os.path.exists(mm_path):
        p2 = htable.rows[0].cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Increased height to 0.6 for better visibility
        p2.add_run().add_picture(mm_path, height=Inches(0.6))

    # --- PAGE 1: COVER PAGE ---
    for _ in range(3): doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(iflow_name)
    run.font.size, run.font.bold = Pt(28), True
    run.font.color.rgb = RGBColor(31, 56, 100)
    
    doc.add_paragraph("Technical Specification Document").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    # AUTHOR TABLE (Restored as per request)
    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.style = 'Table Grid'
    for i, (l, v) in enumerate([("Author:", author), ("Date:", date), ("Version:", "1.0")]):
        meta.rows[i].cells[0].text = l
        meta.rows[i].cells[1].text = v
    doc.add_page_break()

    # --- PAGE 2: TABLE OF CONTENTS (FIXED SPACING) ---
    doc.add_heading("Table of Contents", level=1)
    items = [
        ("1. Introduction", 0), ("1.1 Purpose", 1), ("1.2 Scope", 1),
        ("2. Integration Overview", 0), ("2.1 Integration Architecture", 1), ("2.2 Integration Components", 1),
        ("3. Integration Scenarios", 0), ("3.1 Scenario Description", 1), ("3.2 Data Flows", 1), ("3.3 Security Requirements", 1),
        ("4. Error Handling and Logging", 0), ("5. Testing Validation", 0), ("6. Reference Documents", 0)
    ]
    for text, ind in items:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(0) # Removes gaps between lines
        if ind == 1: p.paragraph_format.left_indent = Inches(0.3)
    doc.add_page_break()

    # --- PAGE 3+: CONTENT MAPPING ---
    content_map = [
        ("1.", "Introduction", 1, None),
        ("1.1", "Purpose", 2, "purpose"),
        ("1.2", "Scope", 2, "scope"),
        ("2.", "Integration Overview", 1, None),
        ("2.1", "Integration Architecture", 2, "architecture"),
        ("2.2", "Integration Components", 2, "components"),
        ("3.", "Integration Scenarios", 1, None),
        ("3.1", "Scenario Description", 2, "scenario_desc"),
        ("3.2", "Data Flows", 2, "data_flows"),
        ("3.3", "Security Requirements", 2, "security"),
        ("4.", "Error Handling and Logging", 1, "error_handling"),
        ("5.", "Testing Validation", 1, "testing"),
        ("6.", "Reference Documents", 1, "references")
    ]

    for num, title, lvl, key in content_map:
        add_numbered_heading(doc, title, lvl, num)
        if key:
            val = data.get(key, f"Refer to technical design for {title}.")
            p = doc.add_paragraph(str(val))
            p.paragraph_format.space_after = Pt(6)

    doc.save(output_path)

if __name__ == "__main__":
    build_docx(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4], sys.argv[5])
